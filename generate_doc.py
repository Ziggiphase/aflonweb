import subprocess
import sys
import os

# Ensure python-docx is installed quietly
try:
    import docx
except ImportError:
    print("Installing required package 'python-docx'...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
    import docx

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

document = Document()

# Title
title = document.add_heading('Strategic Recommendation: Web Hosting Infrastructure', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

document.add_paragraph()

# Intro
p = document.add_paragraph()
p.add_run('Objective: ').bold = True
p.add_run('To outline the most efficient, secure, and profitable approach for upgrading AflonWeb’s hosting infrastructure to support automated web hosting and domain sales.')

# Section 1
document.add_heading('1. The Need for an Upgrade', level=1)
p1 = document.add_paragraph('Currently, AflonWeb is operating on a standard "Shared Hosting" plan (WP Power). While fine for a single website, it is not designed to host multiple clients securely or automate sales.')
p1_2 = document.add_paragraph('To operate like a professional hosting company (similar to WhoGoHost or Upperlink), we must upgrade to a ')
p1_2.add_run('Reseller Hosting Plan').bold = True
p1_2.add_run('. A Reseller plan provides the necessary control panel tools (WHM) to automatically generate isolated, secure accounts for every new customer the moment they pay on our website.')

# Section 2
document.add_heading('2. Managed Reseller (InMotion) vs. Raw Cloud (Kamatera / DigitalOcean)', level=1)
p2 = document.add_paragraph('While building our own servers from scratch using raw cloud providers like Kamatera or DigitalOcean is an option, it is not recommended for our current stage of growth. Here is why we should stay with a Managed Reseller plan via InMotion:')

# Bullet points
document.add_paragraph('Zero Maintenance Overhead: With InMotion, they act as the "landlord." They handle server security, hardware failures, and network outages. With raw cloud (Kamatera), we are 100% responsible for fixing the server if it crashes at 3:00 AM.', style='List Bullet')
document.add_paragraph('Hidden Software Costs: Raw cloud servers appear cheap initially (e.g., $15/month). However, they require us to purchase our own software licenses (cPanel/WHM) to manage clients, which adds $40+ per month to the cost. InMotion bundles these expensive licenses for free.', style='List Bullet')
document.add_paragraph('Focus on Revenue, Not IT: At this stage, our team\'s energy is best spent acquiring clients, closing sales, and building websites. Managing complex Linux servers distracts from revenue generation.', style='List Bullet')

# Section 3
document.add_heading('3. The Recommended Action Plan', level=1)
document.add_paragraph('Step 1: Upgrade our current InMotion plan to a Reseller Plan.', style='List Number')
document.add_paragraph('Step 2: Implement billing automation software (like WHMCS) to connect our website to InMotion and our NiRA domain account.', style='List Number')
document.add_paragraph('Step 3: Focus aggressively on sales and marketing.', style='List Number')

document.add_paragraph()
p_final = document.add_paragraph()
p_final.add_run('Conclusion: ').bold = True
p_final.add_run('The Managed Reseller route allows us to look and function exactly like a massive enterprise hosting company without the expensive overhead or technical risks of running our own server hardware. We can seamlessly migrate to our own custom cloud infrastructure in the future once we cross 100+ active hosting clients.')

# Save the document
save_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'AflonWeb_Hosting_Strategy.docx')
document.save(save_path)
print(f"Document successfully created at: {save_path}")
